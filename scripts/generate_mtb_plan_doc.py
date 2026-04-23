# -*- coding: utf-8 -*-
from __future__ import annotations

from datetime import date
from pathlib import Path
import json

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "output" / "doc"
OUT_PATH = OUT_DIR / "mtb_galaxy_plan.docx"

OPENAPI = json.loads((ROOT / "apps" / "api" / "openapi.json").read_text(encoding="utf-8"))
WEB_PACKAGE = json.loads((ROOT / "apps" / "web" / "package.json").read_text(encoding="utf-8"))

API_ROWS = [
    ("POST", "/auth/demo-login", "Создание или обновление демо-сессии пользователя."),
    ("GET", "/galaxy/profile", "Получение профиля галактики, планет, бустеров и прогресса."),
    ("POST", "/events/ingest", "Прием банковского события и запуск движка наград."),
    ("GET", "/quests", "Получение текущих квестов пользователя."),
    ("POST", "/quests/{quest_id}/claim", "Выдача награды за выполненный квест."),
    ("GET", "/rewards/ledger", "Журнал начислений и наград."),
    ("POST", "/games/runs", "Сохранение результата мини-игры."),
    ("GET", "/games/summary", "Сводка по мини-играм и прогрессу."),
    ("GET", "/friends", "Получение списков друзей, входящих и исходящих приглашений."),
    ("POST", "/friends/invite", "Отправка приглашения в друзья вручную, по QR или реферальному сценарию."),
    ("POST", "/friends/accept", "Подтверждение входящего приглашения и фиксация дружеской связи."),
    ("GET", "/friends/activity", "Лента активности друзей и социального контура."),
    ("GET", "/qr/me", "Получение собственного QR payload и CTA для дальнейшего сценария."),
    ("POST", "/qr/resolve", "Разбор QR payload и определение действия: друзья, реферал или AI."),
    ("GET", "/assistant/context", "Сбор контекста AI-помощника по друзьям, квестам и QR-сценариям."),
    ("POST", "/assistant/chat", "Диалог с AI-помощником с учетом сообщения пользователя и QR payload."),
    ("POST", "/referrals/invite", "Создание приглашения друга."),
    ("GET", "/leaderboard", "Получение лидерборда игроков."),
    ("GET", "/admin/kpi", "Витрина KPI для административного контура."),
    ("POST", "/admin/simulate", "Симуляция банковского события из админ-панели."),
    ("GET", "/admin/risk", "Просмотр риск-флагов и ожидающих наград."),
    ("GET", "/admin/stream", "SSE-поток обновления административных метрик."),
]

PAGE_ROWS = [
    ("GalaxyPage", "Главный экран галактики с прогрессом, планетами и действиями пользователя."),
    ("PlanetsMapScreen / PlanetDetailScreen", "Планетарная карта и детальный просмотр каждой планеты."),
    ("GamesPage / GameScreen", "Запуск мини-игр и управление доступностью игровых сценариев."),
    ("QuestsPage", "Отображение квестов и получение наград."),
    ("RewardsPage", "История наград, звездной пыли и бонусов."),
    ("FriendsPage (/app/friends)", "Полноценный social MVP: списки друзей, инвайты, входящие заявки и переходы в QR-сценарий."),
    ("QrPage (/app/qr)", "Экран собственного QR и резолвера payload для маршрутов друзей, рефералов и AI-помощника."),
    ("AiPage (/app/ai)", "AI-ассистент с контекстом по друзьям, квестам, QR и быстрыми сценариями навигации."),
    ("ReferralsPage / LeaderboardScreen", "Реферальный сценарий и турнирная таблица игроков."),
    ("AdminKpiPage / AdminSimulatorPage / AdminRiskPage", "Административный контур: KPI, симулятор событий и риск-мониторинг."),
]

PLAN_ROWS = [
    ("FastAPI + база данных", "Выполнено", "Реализован backend на FastAPI; по умолчанию используется SQLite, опционально поддерживаются PostgreSQL и Redis."),
    ("Планетарная карта", "Выполнено", "Сценарий реализован через GalaxyPage и экран карты планет."),
    ("Детализация планет/созвездий", "Выполнено", "Есть отдельный экран планеты с прогрессом, зданиями и играми."),
    ("Социальный модуль / друзья", "Выполнено", "Есть FriendsPage по маршруту /app/friends и backend-методы /friends, /friends/invite, /friends/accept, /friends/activity."),
    ("Профиль пользователя", "Выполнено", "Профиль и сессия доступны через демо-логин, состояние хранится на клиенте и в API."),
    ("Игровые механики и подарки", "Выполнено с расширением", "Добавлены квесты, журнал наград, звездная пыль и несколько мини-игр."),
    ("AI и QR экраны из исходного плана", "Выполнено", "Реализованы QrPage (/app/qr) и AiPage (/app/ai), а backend включает /qr/me, /qr/resolve, /assistant/context и /assistant/chat."),
]

PROJECT_TREE = """MTB/
├── apps/
│   ├── api/
│   │   ├── app/
│   │   │   ├── main.py
│   │   │   ├── api.py
│   │   │   ├── db.py
│   │   │   ├── models.py
│   │   │   ├── schemas.py
│   │   │   ├── engine.py
│   │   │   ├── quests.py
│   │   │   ├── analytics.py
│   │   │   └── worker.py
│   │   ├── tests/
│   │   └── openapi.json
│   └── web/
│       ├── src/app/router.tsx
│       ├── src/pages/
│       │   ├── GalaxyPage.tsx
│       │   ├── GamesPage.tsx
│       │   ├── QuestsPage.tsx
│       │   ├── RewardsPage.tsx
│       │   ├── FriendsPage.tsx
│       │   ├── QrPage.tsx
│       │   ├── AiPage.tsx
│       │   ├── ReferralsPage.tsx
│       │   └── Admin*.tsx
│       ├── src/features/planets/screens/
│       └── src/lib/api.ts
├── packages/contracts/src/generated/api.ts
├── scripts/serve-web.mjs
├── docker-compose.yml
└── README.md"""

LAUNCH_BLOCK = r"""npm install
python -m venv .venv
.venv\Scripts\activate
pip install -e apps/api[dev]
npm run contracts:generate
npm run dev"""

CHECK_BLOCK = r"""npm run test:api
npm run test:web
npm run build:web
npm run contracts:generate"""

REQS = [
    "Node.js 20 или новее.",
    "Python 3.12 или новее.",
    "Windows PowerShell, Git Bash или другая оболочка с доступом к npm и Python.",
]

IMPLEMENTED = [
    "Игровой web-клиент на React 19, Vite 6, React Router 7, Zustand, TanStack Query и Framer Motion.",
    "Backend на FastAPI с SQLAlchemy, Pydantic и локальной SQLite-базой по умолчанию.",
    "Контракты API в пакете packages/contracts, синхронизированные с OpenAPI.",
    "Планеты, квесты, награды, лидерборд, реферальная механика и административный контур.",
    "Реально работающие social/QR/AI модули: маршруты /app/friends, /app/qr, /app/ai и соответствующие backend endpoint-ы /friends, /qr/*, /assistant/*.",
    "Набор мини-игр: Halva Snake, Credit Shield Reactor, Social Ring Signal и дополнительные игровые страницы MVP.",
]

GAME_BULLETS = [
    "ORBIT_COMMERCE: сценарии покупок, прогресс планеты и связанные игровые активности.",
    "CREDIT_SHIELD: дисциплина кредитования, игровые проверки и защитные механики.",
    "SOCIAL_RING: реферальная логика, приглашения и социальный прогресс.",
    "Система наград начисляет звездную пыль, бустеры и контейнеры; результаты попадают в журнал наград.",
]


def set_font(run, name: str = "Times New Roman", size: float = 14, bold: bool = False, italic: bool = False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def set_cell_text(cell, text: str, bold: bool = False, size: float = 11) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_default_style(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(14)
    for section in document.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)


def add_paragraph(
    document: Document,
    text: str,
    first_line: float = 1.25,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    after: float = 6,
    before: float = 0,
    spacing: float = 1.5,
    bold: bool = False,
    italic: bool = False,
):
    paragraph = document.add_paragraph()
    paragraph.alignment = align
    if first_line:
        paragraph.paragraph_format.first_line_indent = Cm(first_line)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.line_spacing = spacing
    run = paragraph.add_run(text)
    set_font(run, bold=bold, italic=italic)
    return paragraph


def add_heading(document: Document, text: str, level: int = 1):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.2
    run = paragraph.add_run(text)
    set_font(run, size=15 if level == 1 else 14, bold=True)
    return paragraph


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.3
        run = paragraph.add_run(item)
        set_font(run)


def add_code_block(document: Document, text: str) -> None:
    for line in text.splitlines():
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.left_indent = Cm(0.75)
        paragraph.paragraph_format.right_indent = Cm(0.3)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run(line)
        set_font(run, name="Consolas", size=9.5)


def add_table(document: Document, headers: list[str], rows: list[tuple[str, ...]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, size=11)
        shade_cell(table.rows[0].cells[index], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], str(value), size=10.5)
    document.add_paragraph("")


def build_document() -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_default_style(doc)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(12)
    paragraph.paragraph_format.space_before = Pt(12)
    run = paragraph.add_run("MTB Galaxy")
    set_font(run, size=18, bold=True)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(16)
    run = paragraph.add_run("Реализация проекта по исходному плану в формате Word")
    set_font(run, size=14, bold=True)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(20)
    run = paragraph.add_run(f"Актуализировано по состоянию на {date.today().strftime('%d.%m.%Y')}")
    set_font(run, size=12, italic=True)

    add_paragraph(
        doc,
        "Документ подготовлен на основе исходного Word-плана и текущей рабочей версии проекта MTB Galaxy в каталоге D:\\MTB. Ниже зафиксированы назначение системы, архитектура, соответствие исходному плану, основные модули, API и порядок запуска приложения.",
    )

    add_heading(doc, "1. Цель и исходный план")
    add_paragraph(
        doc,
        "Исходный документ задавал разработку приложения MTB Galaxy как геймифицированного банковского продукта с планетарной картой, социальными механиками, наградами и backend-частью на FastAPI. Базовый план ориентировался на мобильный стек React Native + FastAPI и предусматривал экраны галактики, деталей планет, друзей, AI, QR и профиля.",
    )
    add_paragraph(
        doc,
        "Текущая рабочая версия проекта реализует ту же бизнес-идею в формате web-MVP: пользователь прокачивает планеты, выполняет квесты, получает награды, запускает мини-игры и взаимодействует с административным контуром. Таким образом, исходный план не просто повторен, а расширен и приведен к демонстрационно-боевому виду.",
    )

    add_heading(doc, "2. Что реализовано в текущей версии")
    add_bullets(doc, IMPLEMENTED)

    add_heading(doc, "3. Соответствие исходному плану")
    add_paragraph(
        doc,
        "Ниже приведено сопоставление основных пунктов исходного Word-плана и фактической реализации в репозитории. Это позволяет показать, какие элементы полностью сохранены и какие были расширены относительно базовой постановки.",
    )
    add_table(doc, ["Пункт плана", "Статус", "Фактическая реализация"], PLAN_ROWS)

    add_heading(doc, "4. Архитектура проекта")
    add_paragraph(
        doc,
        "Проект построен как монорепозиторий. Пользовательский интерфейс расположен в apps/web, серверная логика и API находятся в apps/api, а типы и контракты вынесены в отдельный пакет packages/contracts. Такой подход упрощает синхронизацию схем, ускоряет разработку и уменьшает риск расхождения между фронтендом и бэкендом.",
    )
    add_code_block(doc, PROJECT_TREE)
    add_paragraph(
        doc,
        "Корневой package.json содержит сценарии для локальной разработки, одновременного запуска API и web-клиента, генерации контрактов из OpenAPI и запуска тестов. За счет этого проект можно разворачивать как единый стенд для демонстрации и дальнейшей доработки.",
    )

    add_heading(doc, "5. Backend: FastAPI, модели и движок наград")
    add_paragraph(
        doc,
        "Серверная часть разработана на FastAPI. Точка входа находится в apps/api/app/main.py: приложение поднимает экземпляр FastAPI, настраивает CORS, создает таблицы и выполняет сидирование демонстрационных данных при старте. Маршруты вынесены в модуль api.py, бизнес-логика прогресса и начислений разделена по модулям engine.py, quests.py, analytics.py и worker.py.",
    )
    add_paragraph(
        doc,
        "Для хранения данных по умолчанию используется SQLite, однако конфигурация предусматривает расширение до PostgreSQL. Модели и схемы построены на SQLAlchemy и Pydantic. Это позволяет обслуживать как локальный демо-стенд, так и более боевую конфигурацию без полной перестройки архитектуры.",
    )
    add_bullets(
        doc,
        [
            "FastAPI обеспечивает HTTP API и OpenAPI-описание интерфейса.",
            "SQLAlchemy отвечает за модели, сессии и доступ к данным.",
            "Pydantic используется для валидации входных и выходных схем.",
            "Seed-логика создает демо-пользователя и стартовые квесты при запуске.",
            "Движок наград обрабатывает игровые и банковские события, влияющие на прогресс пользователя.",
        ],
    )

    add_heading(doc, "6. Frontend: web-клиент MTB Galaxy")
    add_paragraph(
        doc,
        f"Клиентская часть написана на React {WEB_PACKAGE['dependencies']['react'].lstrip('^')} с использованием Vite {WEB_PACKAGE['devDependencies']['vite'].lstrip('^')}, React Router {WEB_PACKAGE['dependencies']['react-router-dom'].lstrip('^')}, Zustand и TanStack Query. Маршрутизация реализована в apps/web/src/app/router.tsx: приложение разделено на пользовательский контур и административные страницы.",
    )
    add_paragraph(
        doc,
        "Главный экран GalaxyPage является центром сценария: на нем показаны планеты, игровые действия, статистика, прогресс и ссылки на остальные разделы. Дополнительно реализованы карты планет, отдельные игровые страницы, квесты, награды, реферальный раздел и админский блок. Интерфейс построен как единое SPA-приложение, что ускоряет демонстрацию и упрощает перенос в разные окружения.",
    )
    add_paragraph(
        doc,
        "Отдельно от главного контура в web-клиенте присутствуют маршруты /app/friends, /app/qr и /app/ai. Они ведут соответственно на FriendsPage, QrPage и AiPage и закрывают именно те сценарии, которые были заложены в исходном плане: дружеские связи, QR-переходы и AI-помощник.",
    )
    add_table(doc, ["Ключевой экран/модуль", "Назначение"], PAGE_ROWS)

    add_heading(doc, "7. Игровая логика и банковые механики")
    add_paragraph(
        doc,
        "В проекте выделены три основные планеты прогресса, отражающие ключевые банковские сценарии пользователя. Каждая планета имеет собственную логику действий, условия разблокировки и связанную игровую активность. Дополнительно в MVP присутствуют мини-игры и система наград, поддерживающие вовлечение пользователя.",
    )
    add_bullets(doc, GAME_BULLETS)
    add_paragraph(
        doc,
        "Важной частью реализации является административный контур. Он позволяет демонстрировать продуктовые KPI, моделировать события и отслеживать риск-флаги. За счет этого проект подходит не только как пользовательский прототип, но и как демо-стенд для внутренних команд и презентаций.",
    )

    add_heading(doc, "8. Основные API-методы")
    add_paragraph(
        doc,
        f"OpenAPI-описание проекта содержит {sum(len(value) for value in OPENAPI.get('paths', {}).values())} основных endpoint-ов. Они покрывают авторизацию демо-пользователя, профиль галактики, события, квесты, награды, мини-игры, социальный модуль друзей, QR-механику, AI-помощника, рефералов, лидерборд и административный функционал.",
    )
    add_paragraph(
        doc,
        "Для social/QR/AI контура в репозитории уже опубликованы и используются endpoint-ы /friends, /friends/invite, /friends/accept, /friends/activity, /qr/me, /qr/resolve, /assistant/context и /assistant/chat. Это не заглушки и не адаптационные заметки, а часть текущего рабочего API и маршрутизации web-клиента.",
    )
    add_table(doc, ["Метод", "Путь", "Назначение"], API_ROWS)

    add_heading(doc, "9. Требования и запуск проекта")
    add_paragraph(
        doc,
        "Для локального запуска необходимо подготовить стандартную среду разработки и установить зависимости как для JavaScript-части, так и для Python API. Базовые требования и рекомендуемый сценарий запуска соответствуют README проекта.",
    )
    add_bullets(doc, REQS)
    add_paragraph(doc, "Рекомендуемая последовательность команд для старта стенда:")
    add_code_block(doc, LAUNCH_BLOCK)
    add_paragraph(
        doc,
        "После запуска приложение открывается по адресу http://localhost:5173, а API работает на порту 8000. Для синхронизации контракта фронтенда с FastAPI используется отдельная команда генерации OpenAPI-типов.",
    )

    add_heading(doc, "10. Проверка качества")
    add_paragraph(
        doc,
        "Для проверки целостности проекта предусмотрены автоматические команды тестирования и сборки. Их следует выполнять перед демонстрацией, публикацией или передачей проекта. Набор проверок покрывает backend, frontend, сборку web-клиента и генерацию контрактов.",
    )
    add_code_block(doc, CHECK_BLOCK)
    add_paragraph(
        doc,
        "Дополнительно рекомендуется ручной прогон ключевых экранов на нескольких ширинах экрана, а также проверка игровых сценариев, квестов, лидерборда и административной панели. Такой порядок позволяет быстро убедиться, что проект соответствует ожиданиям исходного плана и корректно показывает расширенный функционал MVP.",
    )

    add_heading(doc, "11. Вывод")
    add_paragraph(
        doc,
        "По итогам анализа можно сделать вывод, что MTB Galaxy реализован как полноценный демонстрационный продукт на базе FastAPI и современного React web-стека. По отношению к исходному Word-плану проект не только сохраняет ключевые идеи планетарной карты, прогресса, наград и социального сценария, но и расширяет их за счет квестов, административного контура, контрактов API и нескольких мини-игр.",
    )
    add_paragraph(
        doc,
        "Итоговая версия документа подходит как для сдачи, внутренней презентации или сопровождения проекта: она фиксирует архитектуру, текущее состояние реализации и порядок запуска без необходимости показывать сырой набор файлов и черновые фрагменты кода из исходного плана.",
    )

    doc.save(str(OUT_PATH))
    return OUT_PATH


if __name__ == "__main__":
    print(build_document())
